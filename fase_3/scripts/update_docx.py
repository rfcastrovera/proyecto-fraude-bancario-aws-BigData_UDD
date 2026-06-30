"""Update Propuesta_Arquitectonica.docx with latest changes."""
from docx import Document
from docx.shared import Pt
import re

SRC = "/Users/adrianespinoza/Documents/Documentos - MacBook Air de Adrian/Ciencia de datos UDD/Cloud/fase_3/Propuesta_Arquitectonica_Fase3_Fraude_V1_26062026.docx"
DST = "/Users/adrianespinoza/Documents/Documentos - MacBook Air de Adrian/Ciencia de datos UDD/Cloud/fase_3/Propuesta_Arquitectonica_Fase3_Fraude_V1_30062026.docx"

doc = Document(SRC)
paras = doc.paragraphs

# Helper: set text keeping original formatting
def set_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

# Map paragraph index → new text (from markdown content)
updates = {
    20: (
        "- **Orquestador (EventBridge Scheduler + Lambda)**: Controla el flujo de trabajo en secuencia, "
        "iniciando el job Glue de ingesta, esperando su finalización mediante polling, y luego disparando el job de transformación.\n"
        "- **Procesamiento (AWS Glue/Spark)**: Ejecuta la lógica de negocio, consumiendo recursos DPUs dinámicos.\n"
        "- **Almacenamiento (Medallion)**: S3/Delta Lake actúa como persistencia ACID, permitiendo trazabilidad.\n"
        "- **Consumo (Athena)**: Capa SQL serverless orientada a la toma de decisiones."
    ),
    27: (
        "El sistema procesa un dataset de 10 millones de filas sintéticas (~2.3 GB en JSON, ~143 MB en Parquet comprimido) "
        "generadas con Faker. El desafío técnico principal fue la gestión del Shuffle en los Joins durante la agregación, "
        "superando las limitaciones de un esquema plano Parquet."
    ),
    29: (
        "La arquitectura se fundamenta en el patrón Medallion:\n"
        "- Bronze: Lectura de JSON desde S3, adición de metadatos de linaje (_ingestion_ts, _source_file), escritura como Parquet. 10M filas, ~143 MB comprimido.\n"
        "- Silver: Filtro de calidad (solo montos positivos), enmascaramiento PII de tarjetas (4532-XXXX-XXXX-0367), escritura como Delta Lake particionado por transaction_date. 9.8M filas resultantes.\n"
        "- Gold: Agregación diaria por usuario (tx_count, daily_total, is_high_risk), escritura como Delta Lake. 1.8M filas, 10K usuarios únicos, 50.24% clasificados como alto riesgo.\n\n"
        "El pipeline completo se ejecuta en ~4.5 minutos con 5 workers G.1X (5 DPU) en AWS Glue 5.0. "
        "La orquestación diaria a las 03:00 UTC se realiza mediante EventBridge Scheduler + Lambda, "
        "que inicia los jobs Glue secuencialmente con espera por polling."
    ),
    31: (
        "Utilizando las métricas obtenidas vía Spark UI durante la fase 2 con dataset controlado de 50K registros, "
        "logramos las siguientes mejoras:\n"
        "- Tuning de Particionamiento: spark.sql.shuffle.partitions de 200 → 32. Reducción del 88% en latencia de tareas.\n"
        "- Join Strategies: BroadcastHashJoin forzado para tablas pequeñas. Reducción del 70% en tráfico de red entre Executors.\n"
        "- Caching Estratégico: .cache() solo tras filtros críticos, .unpersist() al final del DAG. Optimización de 45% de memoria heap."
    ),
}

# Insert new paragraphs after index 24 (before "3. Informe Técnico")
# We need to add the orchestration flow section
# Since inserting is complex, we'll add it as a replacement of the empty paragraph after 24

# Actually, let's find the right place and add content
# Paragraph 25 is "3. Informe Técnico" — we'll insert before it

# Find the element in the document body
body = doc.element.body
# Find the paragraph element for "3. Informe Técnico" (index 25 in paras list)
p_elements = body.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
target_p = p_elements[25]  # "3. Informe Técnico"

# Create new paragraphs before it
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# Helper to add a new paragraph before a reference element
def insert_paragraph_before(ref_element, text, style='Normal', bold=False, size=None):
    new_p = OxmlElement('w:p')
    # Copy style from reference
    pPr = OxmlElement('w:pPr')
    if style:
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style)
        pPr.append(pStyle)
    new_p.append(pPr)
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))  # half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    
    ref_element.addprevious(new_p)
    return new_p

# Insert Section 3: Flujo de Orquestacion
insert_paragraph_before(target_p, "3. Flujo de Orquestación", bold=True, size=14)
insert_paragraph_before(target_p, "EventBridge Scheduler (cron: 0 3 * * ? *) ── Lambda (glue-fraude-orchestrator) ── StartGlueJob(ingestion) ── poll ── StartGlueJob(transform) ── poll ── Return {status, metrics}")
insert_paragraph_before(target_p, "Costo: ~$1/mo (Scheduler) + Lambda (tier gratuito). Alternativa MWAA: DAG preparado en dags/fraude_mwaa_dag.py.")

# Now update specific paragraphs
for idx, text in updates.items():
    set_text(paras[idx], text)

# Update ADR-003 (MWAA → EventBridge)
for idx, para in enumerate(paras):
    if "ADR-003: Orquestación mediante Airflow (MWAA)" in para.text:
        set_text(para, "ADR-003: Orquestación mediante EventBridge Scheduler + Lambda")
    elif para.text.startswith("Contexto:") and idx > 50 and "pipeline se volvió modular" in para.text:
        set_text(para, "Contexto: El pipeline se volvió modular (ingestion.py, transformations.py), requiriendo gestión de errores y ejecución secuencial.")
    elif para.text.startswith("Decisión:") and idx > 50 and "AWS MWAA" in para.text:
        set_text(para, "Decisión: EventBridge Scheduler + Lambda (alternativa a MWAA no disponible en la cuenta).")
    elif para.text.startswith("Alternativas:") and idx > 50 and "Step Functions" in para.text:
        set_text(para, "Alternativas: MWAA (requiere suscripción no disponible), Step Functions (alternativa válida), Airflow local (solo desarrollo).")
    elif para.text.startswith("Consecuencias:") and idx > 50 and "Permite gestionar" in para.text:
        set_text(para, "Consecuencias: Sin costo fijo mensual, despliegue simple, polling desde Lambda para dependencias entre jobs. Se mantiene dags/fraude_mwaa_dag.py para migración futura a MWAA.")

# Fix "15M+" in ADR-001 context
for para in paras:
    if "Necesidad de procesar datasets de 15M+" in para.text:
        para.runs[0].text = "Contexto: Necesidad de procesar datasets de 10M+ de filas con mínima carga operativa."
        break

doc.save(DST)
print(f"Saved: {DST}")
PYEOF