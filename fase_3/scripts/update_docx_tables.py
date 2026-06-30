"""Add tables to .docx and clean up duplicate paragraphs."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = "/Users/adrianespinoza/Documents/Documentos - MacBook Air de Adrian/Ciencia de datos UDD/Cloud/fase_3/Propuesta_Arquitectonica_Fase3_Fraude_V1_30062026.docx"
DST = "/Users/adrianespinoza/Documents/Documentos - MacBook Air de Adrian/Ciencia de datos UDD/Cloud/fase_3/Propuesta_Arquitectonica_Fase3_Fraude_V1_30062026.docx"

doc = Document(SRC)
body = doc.element.body
paras = doc.paragraphs

def insert_table_after(ref_para, headers, rows, title=None):
    """Insert a formatted table after a given paragraph."""
    ref_element = ref_para._element
    
    if title:
        title_p = OxmlElement('w:p')
        title_r = OxmlElement('w:r')
        title_rPr = OxmlElement('w:rPr')
        title_b = OxmlElement('w:b')
        title_rPr.append(title_b)
        title_r.append(title_rPr)
        title_t = OxmlElement('w:t')
        title_t.text = title
        title_t.set(qn('xml:space'), 'preserve')
        title_r.append(title_t)
        title_p.append(title_r)
        ref_element.addnext(title_p)
        ref_element = title_p
    
    # Add empty paragraph before table
    empty_p = OxmlElement('w:p')
    empty_r = OxmlElement('w:r')
    empty_t = OxmlElement('w:t')
    empty_t.set(qn('xml:space'), 'preserve')
    empty_t.text = ""
    empty_r.append(empty_t)
    empty_p.append(empty_r)
    ref_element.addnext(empty_p)
    
    # Create table
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    
    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Inches(1.8)
    
    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    
    # Move table element to after the empty paragraph
    tbl_element = table._tbl
    ref_element.addnext(tbl_element)
    
    return table


def remove_paragraph(para):
    """Remove a paragraph from the document body."""
    p = para._element
    p.getparent().remove(p)


# === 1. CLEAN UP DUPLICATES ===

# Remove old paragraph [21] "Orquestador (MWAA)"
for p in paras:
    if p.text.strip().startswith("Orquestador (MWAA)"):
        remove_paragraph(p)
        break

# Fix ADR-001 context (had duplicate text)
for p in paras:
    if "Necesidad de procesar datasets de 10M+" in p.text and "Necesidad de procesar datasets de 15M+" in p.text:
        # Remove the duplicate
        for run in p.runs:
            if "Necesidad de procesar datasets de 15M+" in run.text:
                run.text = run.text.replace("Necesidad de procesar datasets de 15M+ de filas con mínima carga operativa.", "")
        break

# === 2. INSERT TABLES ===

# Find paragraphs to insert after
opt_para = None  # after optimization intro → metrics table
arch_para = None  # after architecture description → costs table
limit_para = None  # before limitaciones → athena results

for p in paras:
    t = p.text.strip()
    if t.startswith("Utilizando las métricas obtenidas vía Spark UI"):
        opt_para = p
    elif t.startswith("3.4. Limitaciones"):
        limit_para = p
    elif t.startswith("La arquitectura se fundamenta"):
        arch_para = p

# Table 1: Métricas de Producción
if opt_para:
    insert_table_after(
        opt_para,
        ["Etapa", "Runtime", "Workers", "DPU-Seconds", "Resultado"],
        [
            ["Ingesta Bronze", "109 s (1.8 min)", "5 × G.1X", "546", "✅ SUCCEEDED"],
            ["Transform Silver-Gold", "163 s (2.7 min)", "5 × G.1X", "817", "✅ SUCCEEDED"],
            ["Pipeline total", "272 s (~4.5 min)", "5 DPU", "1,363", "✅"],
        ],
        title="Tabla 1: Métricas de Producción (AWS Glue 5.0 — 10M registros)"
    )

# Table 2: Resultados Athena (before limitaciones)
if limit_para:
    insert_table_after(
        limit_para,
        ["Consulta", "Resultado"],
        [
            ["Top 10 usuarios más activos", "Máx: 20 transacciones/día, $927"],
            ["Usuarios alto riesgo (>12 tx/día)", "7,003 combinaciones usuario-fecha"],
            ["Volumen diario últimos 7 días", "Funcional con date_parse"],
            ["Métricas de riesgo global", "10,000 usuarios, 50.24% high risk"],
            ["Transacciones por moneda", "USD 50%, EUR 20%, CLP 15%, BRL 10%, MXN 5%"],
            ["Total filas por capa", "Silver: 9,800,000 / Gold: 1,801,955"],
        ],
        title="Tabla 2: Resultados Athena"
    )

# Table 3: Costos Operativos (after architecture implementation)
if arch_para:
    insert_table_after(
        arch_para,
        ["Componente", "Cálculo", "Costo/mes"],
        [
            ["Glue (5 DPU × 0.075 h × 30 días)", "5 × 0.075 × 30 = 11.25 DPU-hours", "$4.95"],
            ["S3 Almacenamiento (~11 GB)", "11 GB × $0.023/GB", "$0.25"],
            ["S3 Solicitudes (~500K)", "—", "$0.05"],
            ["EventBridge Scheduler", "1 schedule mensual", "$1.00"],
            ["Lambda (invocaciones + 600s)", "Tier gratuito", "$0.00"],
            ["Athena (~10 TB escaneados)", "Consultas esporádicas", "$5.00"],
            ["Total", "", "~$10.25/mes"],
        ],
        title="Tabla 3: Costos Operativos Estimados"
    )

# Add comparison note below table 3
for p in paras:
    if "3.4. Limitaciones" in p.text:
        # Add paragraph before limitaciones
        ref_elem = p._element
        note_p = OxmlElement('w:p')
        note_r = OxmlElement('w:r')
        note_t = OxmlElement('w:t')
        note_t.text = "Comparativa: MWAA habría agregado ~$38/mes en costo fijo de entorno mw1.micro."
        note_t.set(qn('xml:space'), 'preserve')
        note_r.append(note_t)
        note_p.append(note_r)
        ref_elem.addprevious(note_p)
        break

doc.save(DST)
print(f"✅ Documento actualizado: {DST}")
EOF